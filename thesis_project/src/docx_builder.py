# -*- coding: utf-8 -*-
"""
Word 生成器 —— 按 WORD_SPEC 生成本科毕业论文草案 (.docx)。

已落实的规范：
  - A4、页边距 2.5cm + 装订线 1cm、页眉/页脚距离
  - 正文：宋体/Times New Roman、小四(12pt)、固定行距 20 磅、首行缩进 2 字符、两端对齐
  - 标题 1/2/3：黑体 三号/四号/13pt，段前段后间距，套用内置样式 -> 可自动生成目录
  - 摘要(小二加粗居中) / 关键词 / 英文摘要
  - 目录域(TOC field) 占位，settings 写 updateFields=true，Word 打开时提示更新（或按 F9）
  - 参考文献标题
  - 中文字体通过 w:eastAsia 正确设置（python-docx 默认只设置西文字体）
"""
from __future__ import annotations
import io
import re
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.format_spec import WORD_SPEC as W
from src.organizer import PLACEHOLDER, iter_node_blocks

_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ---------------------------------------------------------------------------
#  底层：字体（含中文 eastAsia）与行距
# ---------------------------------------------------------------------------
def _set_run_font(run, font_cn, font_en, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_en
    # 西文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)
    rfonts.set(qn("w:eastAsia"), font_cn)   # 关键：中文字体


def _set_line_spacing(pf, spec):
    if "line_spacing_pt" in spec:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(spec["line_spacing_pt"])
    elif "line_spacing" in spec:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = spec["line_spacing"]


# ---------------------------------------------------------------------------
#  页面
# ---------------------------------------------------------------------------
def _setup_page(doc):
    sec = doc.sections[0]
    p = W["page"]
    if p.get("size", "A4").upper() != "A4":
        raise ValueError("当前仅支持 A4 页面尺寸")
    landscape = p.get("orientation", "portrait") == "landscape"
    sec.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    sec.page_height = Cm(21.0 if landscape else 29.7)
    sec.page_width = Cm(29.7 if landscape else 21.0)
    sec.top_margin = Cm(p["margin_top_cm"])
    sec.bottom_margin = Cm(p["margin_bottom_cm"])
    sec.left_margin = Cm(p["margin_left_cm"])
    sec.right_margin = Cm(p["margin_right_cm"])
    sec.gutter = Cm(p["gutter_cm"])
    sec.header_distance = Cm(p["header_distance_cm"])
    sec.footer_distance = Cm(p["footer_distance_cm"])


# ---------------------------------------------------------------------------
#  样式：正文 + 标题 1/2/3
# ---------------------------------------------------------------------------
def _setup_styles(doc):
    # 正文 Normal
    b = W["body"]
    normal = doc.styles["Normal"]
    normal.font.name = b["font_en"]
    normal.font.size = Pt(b["size_pt"])
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), b["font_cn"])
    pf = normal.paragraph_format
    pf.alignment = _ALIGN[b["alignment"]]
    _set_line_spacing(pf, b)
    pf.first_line_indent = Pt(b["size_pt"] * b["first_line_indent_chars"])
    pf.space_before = Pt(b["space_before_pt"])
    pf.space_after = Pt(b["space_after_pt"])

    # 标题 1/2/3
    for lvl, h in W["headings"].items():
        style = doc.styles[f"Heading {lvl}"]
        style.font.name = h["font_en"]
        style.font.size = Pt(h["size_pt"])
        style.font.bold = h["bold"]
        style.element.rPr.rFonts.set(qn("w:eastAsia"), h["font_cn"])
        # 标题颜色恢复为黑色（内置样式默认蓝色）
        style.font.color.rgb = None
        hpf = style.paragraph_format
        hpf.alignment = _ALIGN[h["alignment"]]
        hpf.space_before = Pt(h["space_before_pt"])
        hpf.space_after = Pt(h["space_after_pt"])
        _set_line_spacing(hpf, h)


# ---------------------------------------------------------------------------
#  段落辅助
# ---------------------------------------------------------------------------
def _add_para(doc, text, spec, align=None, indent_chars=None, space_before=None,
              space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = _ALIGN[align]
    pf = p.paragraph_format
    _set_line_spacing(pf, spec)
    if indent_chars is not None:
        pf.first_line_indent = Pt(spec["size_pt"] * indent_chars)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, spec.get("font_cn", "宋体"),
                  spec.get("font_en", "Times New Roman"),
                  spec["size_pt"], spec.get("bold", False))
    return p


def _add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
#  TOC 域
# ---------------------------------------------------------------------------
def _add_toc(doc):
    t = W["toc"]
    _add_para(doc, t["title"],
              {"font_cn": "黑体", "font_en": "Times New Roman",
               "size_pt": 16, "bold": True, "line_spacing_pt": 20},
              align="center", space_after=12)
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    levels = min(max(t["levels"], 1), 9)   # 直接采用 spec 值，clamp 到 Word 允许的 1..9
    instr.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t"); hint.text = "【打开文档后按提示更新域，或按 F9】"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, hint, fldEnd):
        run._element.append(el)


def _enable_update_fields(doc):
    """settings.xml 写 updateFields=true：Word 打开时提示一键更新目录/页码域。

    按 ECMA-376 CT_Settings 序列，updateFields 须位于 footnotePr/endnotePr/compat
    等元素之前，故用锚点插入而非直接 append。
    """
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    for tag in ("w:footnotePr", "w:endnotePr", "w:compat"):
        anchor = settings.find(qn(tag))
        if anchor is not None:
            anchor.addprevious(el)
            return
    settings.append(el)


# ---------------------------------------------------------------------------
#  页码：分节（前置部分罗马数字 / 正文阿拉伯数字从 1 起）
# ---------------------------------------------------------------------------
_PG_FMT = {"roman": "lowerRoman", "arabic": "decimal"}   # spec 值 -> w:pgNumType/@w:fmt


def _set_pg_num_type(section, fmt, start=None):
    """在 section 的 w:sectPr 下写入 w:pgNumType（页码格式 / 起始页码）。"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def _write_footer_page_field(footer):
    """向页脚写入居中 PAGE 域。"""
    p = footer.paragraphs[0]
    p.alignment = _ALIGN[W["page_number"]["alignment"]]
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldEnd):
        run._element.append(el)


def _setup_page_numbers(doc):
    pn = W["page_number"]
    front, body = doc.sections[0], doc.sections[1]
    # section 1：前置部分（封面/摘要/目录）罗马数字
    _set_pg_num_type(front, _PG_FMT[pn["front_matter"]])
    _write_footer_page_field(front.footer)
    # section 2：正文与参考文献，阿拉伯数字从 1 重新编号；页脚断开与前节链接
    _set_pg_num_type(body, _PG_FMT[pn["body"]], start=1)
    body.footer.is_linked_to_previous = False
    _write_footer_page_field(body.footer)


# ---------------------------------------------------------------------------
#  主构建
# ---------------------------------------------------------------------------
def build(thesis, out_path):
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    ab = W["abstract"]

    # 关键词数量校验（占位符不警告；仅提示，不阻断生成）
    kws = thesis["keywords"]
    if kws != [PLACEHOLDER] and not (ab["keywords_min"] <= len(kws) <= ab["keywords_max"]):
        print(f"  [警告] 关键词数量 {len(kws)} 不在规范 "
              f"{ab['keywords_min']}-{ab['keywords_max']} 范围，请检查。")

    # ---------- 封面 ----------
    _add_para(doc, "", W["body"], space_before=60)
    _add_para(doc, "本科毕业论文（设计）",
              {"font_cn": "黑体", "font_en": "Times New Roman",
               "size_pt": 26, "bold": True, "line_spacing": 1.5},
              align="center", space_after=40)
    _add_para(doc, thesis["title"],
              {"font_cn": "黑体", "font_en": "Times New Roman",
               "size_pt": 22, "bold": True, "line_spacing": 1.5},
              align="center", space_after=60)
    for label in (f"作　者：{thesis['author']}", "学　院：<请填写>",
                  "专　业：<请填写>", "指导教师：<请填写>", "日　期：<请填写>"):
        _add_para(doc, label,
                  {"font_cn": "宋体", "font_en": "Times New Roman",
                   "size_pt": 14, "line_spacing": 1.5}, align="center")
    _add_page_break(doc)

    # ---------- 中文摘要 ----------
    _add_para(doc, ab["title_cn"],
              {"font_cn": "宋体", "font_en": "Times New Roman",
               "size_pt": ab["title_size_pt"], "bold": ab["title_bold"],
               "line_spacing": 1.5},
              align=ab["title_alignment"], space_before=12, space_after=18)
    _add_para(doc, thesis["abstract"], W["body"], indent_chars=ab.get("first_line_indent_chars", 0))
    kw = ab["keywords_sep"].join(thesis["keywords"])
    _add_para(doc, ab["keywords_label_cn"] + kw,
              {"font_cn": "黑体", "font_en": "Times New Roman",
               "size_pt": 12, "bold": True, "line_spacing_pt": 20},
              space_before=12)
    _add_page_break(doc)

    # ---------- 英文摘要 ----------
    _add_para(doc, ab["title_en"],
              {"font_cn": "Times New Roman", "font_en": "Times New Roman",
               "size_pt": ab["title_en_size_pt"], "bold": True,
               "line_spacing": 1.5},
              align="center", space_before=12, space_after=18)
    _add_para(doc, thesis["abstract_en"], W["body"], indent_chars=ab.get("first_line_indent_chars", 0))
    _add_para(doc, ab["keywords_label_en"] + "; ".join(thesis["keywords_en"]),
              {"font_cn": "Times New Roman", "font_en": "Times New Roman",
               "size_pt": 12, "bold": True, "line_spacing_pt": 20},
              space_before=12)
    _add_page_break(doc)

    # ---------- 目录 ----------
    _add_toc(doc)

    # ---------- 分节：前置部分(section 1) | 正文+参考文献(section 2) ----------
    doc.add_section(WD_SECTION.NEW_PAGE)

    # ---------- 正文章节 ----------
    body_chapters = [c for c in thesis["chapters"]
                     if c.get("section_role", "body") != "appendix"]
    for ci, ch in enumerate(body_chapters, 1):
        counters = {"table": 0, "figure": 0}
        h = doc.add_heading(level=1)
        run = h.add_run(f"第{_cn_num(ci)}章　{ch['title']}")
        _set_run_font(run, W["headings"][1]["font_cn"],
                      W["headings"][1]["font_en"],
                      W["headings"][1]["size_pt"], True)
        _render_content(doc, ch, ci, counters)
        for si, sub in enumerate(ch.get("subs", []), 1):
            h2 = doc.add_heading(level=2)
            run = h2.add_run(f"{ci}.{si}　{sub['title']}")
            _set_run_font(run, W["headings"][2]["font_cn"],
                          W["headings"][2]["font_en"],
                          W["headings"][2]["size_pt"], True)
            _render_content(doc, sub, ci, counters)
            for ti, sub3 in enumerate(sub.get("subs", []), 1):
                h3 = doc.add_heading(level=3)
                run = h3.add_run(f"{ci}.{si}.{ti}　{sub3['title']}")
                _set_run_font(run, W["headings"][3]["font_cn"],
                              W["headings"][3]["font_en"],
                              W["headings"][3]["size_pt"], True)
                _render_content(doc, sub3, ci, counters)

    # ---------- 参考文献 ----------
    _add_page_break(doc)
    h = doc.add_heading(level=1)
    run = h.add_run(W["reference"]["title"])
    _set_run_font(run, W["headings"][1]["font_cn"],
                  W["headings"][1]["font_en"], W["headings"][1]["size_pt"], True)
    for i, ref in enumerate(thesis["references"], 1):
        _add_para(doc, f"[{i}] {ref}",
                  {"font_cn": "宋体", "font_en": "Times New Roman",
                  "size_pt": 10.5, "line_spacing_pt": 20})

    # ---------- 附录 ----------
    appendix_index = 0
    for ch in thesis["chapters"]:
        if ch.get("section_role") != "appendix":
            continue
        appendix_index += 1
        h = doc.add_heading(level=1)
        label = chr(ord("A") + appendix_index - 1)
        title = re.sub(r"^附录\s*[A-Z一二三四五六七八九十]*\s*", "",
                       ch["title"]).strip() or "附录"
        run = h.add_run(f"附录{label}　{title}")
        _set_run_font(run, W["headings"][1]["font_cn"],
                      W["headings"][1]["font_en"],
                      W["headings"][1]["size_pt"], True)
        counters = {"table": 0, "figure": 0}
        _render_content(doc, ch, appendix_index, counters)

    _setup_page_numbers(doc)
    _enable_update_fields(doc)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
#  图 / 表（三线表 + 题注编号 表1-1 / 图1-1）
# ---------------------------------------------------------------------------
_CAPTION_SPEC = {"font_cn": "宋体", "font_en": "Times New Roman",
                 "size_pt": 10.5, "line_spacing_pt": 20}


def _set_three_line_borders(tbl):
    """三线表：顶线/底线 1.5 磅、表头下线 0.75 磅，其余无框线。"""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag, sz in (("top", "12"), ("bottom", "12")):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for tag in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    # CT_TblPr 序列要求 tblBorders 位于 tblLayout/tblCellMar/tblLook 之前，
    # 用锚点插入（同 _enable_update_fields 的做法）。
    for tag in ("w:tblLayout", "w:tblCellMar", "w:tblLook"):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            anchor.addprevious(borders)
            break
    else:
        tblPr.append(borders)
    for tc in tbl.rows[0]._tr.findall(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcB = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "000000")
        tcB.append(bottom)
        tcPr.append(tcB)


def _add_three_line_table(doc, rows, ci, ti):
    spec = W.get("table", {})
    caption = f"{spec.get('prefix', '表')}{ci}-{ti}　<请填写表题>"
    if spec.get("caption_position", "above") == "above":
        _add_para(doc, caption, _CAPTION_SPEC,
                  align=spec.get("caption_align", "center"))
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, row in enumerate(rows):
        for c_i in range(n_cols):
            cell = tbl.rows[r_i].cells[c_i]
            cell.text = row[c_i] if c_i < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, "宋体", "Times New Roman", 10.5,
                                  bold=(r_i == 0))
    _set_three_line_borders(tbl)
    if spec.get("caption_position", "above") != "above":
        _add_para(doc, caption, _CAPTION_SPEC,
                  align=spec.get("caption_align", "center"))


def _add_image(doc, img, ci, fi):
    spec = W.get("figure", {})
    p = doc.add_paragraph()
    p.alignment = _ALIGN["center"]
    try:
        p.add_run().add_picture(io.BytesIO(img["data"]), width=Cm(14))
    except Exception:  # noqa: BLE001 损坏/不支持的图片格式
        # 占位文字写入同一居中段落，仍输出题注以保持编号连续
        run = p.add_run("<图片插入失败，请手工补图>")
        _set_run_font(run, W["body"]["font_cn"], W["body"]["font_en"],
                      W["body"]["size_pt"])
    caption = f"{spec.get('prefix', '图')}{ci}-{fi}　<请填写图题>"
    _add_para(doc, caption, _CAPTION_SPEC,
              align=spec.get("caption_align", "center"))


def _render_content(doc, node, ci, counters):
    """按原始顺序渲染正文与媒体；兼容旧 paras/tables/images 节点。"""
    for block in iter_node_blocks(node):
        kind = block.get("kind")
        if kind == "table":
            counters["table"] += 1
            _add_three_line_table(doc, block.get("rows") or [],
                                  ci, counters["table"])
        elif kind == "image":
            counters["figure"] += 1
            _add_image(doc, block, ci, counters["figure"])
        elif block.get("text"):
            _add_para(doc, block["text"], W["body"], indent_chars=2)


def _cn_num(n):
    """0..99 转中文数字（含 11/19/20/21/99 边界）。调用方传入的 n 从 1 起。"""
    cn_digit = "零一二三四五六七八九"
    if n < 10:
        return cn_digit[n]
    if n == 10:
        return "十"
    if n < 20:
        return "十" + cn_digit[n - 10]
    tens, ones = divmod(n, 10)
    head = cn_digit[tens] + "十"
    return head + (cn_digit[ones] if ones else "")
