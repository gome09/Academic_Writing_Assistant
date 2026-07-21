# -*- coding: utf-8 -*-
import base64
import io

import docx
from src import readers

# 1x1 像素 PNG
_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQ"
    "DwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def test_read_docx_extracts_inline_image(tmp_path):
    d = docx.Document()
    d.add_paragraph("前文段落。")
    d.add_picture(io.BytesIO(_PNG_1PX))
    d.add_paragraph("后文段落。")
    path = str(tmp_path / "img.docx")
    d.save(path)

    doc = readers.read_docx(path)
    imgs = [b for b in doc["blocks"] if b["kind"] == "image"]
    assert len(imgs) == 1
    assert imgs[0]["data"][:4] == b"\x89PNG"
    assert imgs[0]["ext"] == ".png"
    # 图片出现在两段文字之间
    kinds = [b["kind"] for b in doc["blocks"]]
    assert kinds == ["paragraph", "image", "paragraph"]


def test_read_docx_without_images_unchanged(tmp_path):
    d = docx.Document()
    d.add_paragraph("只有文字。")
    path = str(tmp_path / "plain.docx")
    d.save(path)
    doc = readers.read_docx(path)
    assert all(b["kind"] != "image" for b in doc["blocks"])
