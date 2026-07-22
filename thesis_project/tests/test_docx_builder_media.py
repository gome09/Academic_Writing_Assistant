# -*- coding: utf-8 -*-
import base64

import docx
from src import docx_builder

_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQ"
    "DwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def _min_thesis():
    return {"title": "题目", "author": "作者",
            "abstract": "本文摘要。", "abstract_en": "EN abstract.",
            "keywords": ["a", "b", "c"], "keywords_en": ["a", "b", "c"],
            "chapters": [{"title": "绪论", "level": 1, "paras": ["正文。"],
                          "subs": [], "tables": [], "images": []}],
            "auto_skeleton": False, "references": ["某文献[J]. 2024."]}


def test_real_table_rendered_with_caption(tmp_path):
    thesis = _min_thesis()
    thesis["chapters"][0]["tables"] = [[["指标", "数值"], ["准确率", "94%"]]]
    out = str(tmp_path / "o.docx")
    docx_builder.build(thesis, out)

    d = docx.Document(out)
    assert len(d.tables) == 1
    assert d.tables[0].cell(0, 0).text == "指标"
    assert d.tables[0].cell(1, 1).text == "94%"
    texts = [p.text for p in d.paragraphs]
    assert any(t.startswith("表1-1") for t in texts)


def test_image_rendered_with_caption(tmp_path):
    thesis = _min_thesis()
    thesis["chapters"][0]["images"] = [{"data": _PNG_1PX, "ext": ".png"}]
    out = str(tmp_path / "o.docx")
    docx_builder.build(thesis, out)

    d = docx.Document(out)
    assert len(d.inline_shapes) == 1
    texts = [p.text for p in d.paragraphs]
    assert any(t.startswith("图1-1") for t in texts)


def test_broken_image_falls_back_to_placeholder(tmp_path):
    thesis = _min_thesis()
    thesis["chapters"][0]["images"] = [{"data": b"not-an-image", "ext": ".png"}]
    out = str(tmp_path / "o.docx")
    docx_builder.build(thesis, out)          # 不应抛异常
    d = docx.Document(out)
    texts = [p.text for p in d.paragraphs]
    assert any("图片插入失败" in t for t in texts)
    # 占位仍占用编号并输出题注，保持后续图片编号连续
    assert any(t.startswith("图1-1") for t in texts)


def test_table_numbering_resets_per_chapter(tmp_path):
    thesis = _min_thesis()
    thesis["chapters"][0]["tables"] = [[["a", "b"], ["1", "2"]]]
    thesis["chapters"].append({"title": "方法", "level": 1, "paras": ["正文。"],
                               "subs": [], "images": [],
                               "tables": [[["c", "d"], ["3", "4"]]]})
    out = str(tmp_path / "o.docx")
    docx_builder.build(thesis, out)
    d = docx.Document(out)
    texts = [p.text for p in d.paragraphs]
    assert any(t.startswith("表1-1") for t in texts)
    assert any(t.startswith("表2-1") for t in texts)


def test_chapters_without_media_keys_still_build(tmp_path):
    thesis = _min_thesis()
    del thesis["chapters"][0]["tables"], thesis["chapters"][0]["images"]
    docx_builder.build(thesis, str(tmp_path / "o.docx"))   # .get() 容错
