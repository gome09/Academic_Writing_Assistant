# -*- coding: utf-8 -*-
"""页码分节（前置罗马 / 正文阿拉伯从 1 起）、TOC levels 以 spec 为准、关键词数量警告。"""
import docx
from docx.oxml.ns import qn

from src import docx_builder
from config.format_spec import WORD_SPEC as W


def _thesis(keywords=None):
    return {
        "title": "T", "author": "A",
        "abstract": "AB", "abstract_en": "PLACE",
        "keywords": keywords if keywords is not None else ["k1", "k2", "k3"],
        "keywords_en": ["PLACE"],
        "chapters": [{"title": "章1", "level": 1, "paras": ["P"], "subs": []}],
        "auto_skeleton": False, "references": ["示例."],
    }


def _build(tmp_path, keywords=None):
    out = str(tmp_path / "s.docx")
    docx_builder.build(_thesis(keywords), out)
    return out


def _pg_num_type(section):
    return section._sectPr.find(qn("w:pgNumType"))


# ---------------------------------------------------------------------------
#  项 1：页码分节
# ---------------------------------------------------------------------------
def test_docx_has_two_sections(tmp_path):
    doc = docx.Document(_build(tmp_path))
    assert len(doc.sections) >= 2


def test_front_matter_page_number_roman(tmp_path):
    doc = docx.Document(_build(tmp_path))
    pg = _pg_num_type(doc.sections[0])
    assert pg is not None
    assert pg.get(qn("w:fmt")) == "lowerRoman"


def test_body_page_number_decimal_restart(tmp_path):
    doc = docx.Document(_build(tmp_path))
    pg = _pg_num_type(doc.sections[1])
    assert pg is not None
    assert pg.get(qn("w:fmt")) == "decimal"
    assert pg.get(qn("w:start")) == "1"


def test_body_footer_unlinked_with_page_field(tmp_path):
    doc = docx.Document(_build(tmp_path))
    footer = doc.sections[1].footer
    assert footer.is_linked_to_previous is False
    xml = footer._element.xml
    assert "PAGE" in xml and "fldChar" in xml


def test_front_footer_has_page_field(tmp_path):
    doc = docx.Document(_build(tmp_path))
    xml = doc.sections[0].footer._element.xml
    assert "PAGE" in xml and "fldChar" in xml


def test_first_chapter_heading_in_section2(tmp_path):
    doc = docx.Document(_build(tmp_path))
    # 分节符落在第一章标题之前：标题段之前应存在段内 sectPr（section 1 的结尾）
    body_children = list(doc.element.body)
    heading_idx = None
    sect_break_idx = None
    for i, el in enumerate(body_children):
        if el.tag == qn("w:p"):
            ppr = el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None \
                    and sect_break_idx is None:
                sect_break_idx = i
            style = None
            if ppr is not None:
                pstyle = ppr.find(qn("w:pStyle"))
                if pstyle is not None:
                    style = pstyle.get(qn("w:val"))
            if style in ("Heading1", "1") and heading_idx is None:
                heading_idx = i
    assert sect_break_idx is not None, "缺少分节符（段内 sectPr）"
    assert heading_idx is not None, "缺少一级标题"
    assert sect_break_idx < heading_idx, "第一章标题应位于 section 2"


def test_section2_geometry_matches_section1(tmp_path):
    doc = docx.Document(_build(tmp_path))
    s1, s2 = doc.sections[0], doc.sections[1]
    for attr in ("page_width", "page_height", "top_margin", "bottom_margin",
                 "left_margin", "right_margin", "gutter",
                 "header_distance", "footer_distance"):
        assert getattr(s1, attr) == getattr(s2, attr), attr


# ---------------------------------------------------------------------------
#  项 2：TOC levels 以 spec 为准
# ---------------------------------------------------------------------------
def test_toc_spec_levels_is_three():
    assert W["toc"]["levels"] == 3


def test_toc_field_uses_spec_levels(tmp_path, monkeypatch):
    # spec 值直接决定 TOC 域，不与 headings 数量耦合
    import re
    import zipfile
    monkeypatch.setitem(W["toc"], "levels", 2)
    out = _build(tmp_path)
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    m = re.search(r'TOC \\o "1-(\d+)"', xml)
    assert m and int(m.group(1)) == 2


# ---------------------------------------------------------------------------
#  项 3：关键词数量校验
# ---------------------------------------------------------------------------
def test_keywords_too_few_warns(tmp_path, capsys):
    _build(tmp_path, keywords=["k1", "k2"])
    assert "关键词数量 2 不在规范 3-5 范围" in capsys.readouterr().out


def test_keywords_too_many_warns(tmp_path, capsys):
    _build(tmp_path, keywords=["k1", "k2", "k3", "k4", "k5", "k6"])
    assert "关键词数量 6 不在规范 3-5 范围" in capsys.readouterr().out


def test_keywords_in_range_no_warning(tmp_path, capsys):
    _build(tmp_path, keywords=["k1", "k2", "k3"])
    assert "关键词数量" not in capsys.readouterr().out


def test_keywords_placeholder_no_warning(tmp_path, capsys):
    _build(tmp_path, keywords=["<请填写>"])
    assert "关键词数量" not in capsys.readouterr().out
