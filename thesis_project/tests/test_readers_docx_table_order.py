# -*- coding: utf-8 -*-
"""read_docx 按文档 body 顺序混排段落与表格（表格不再全部堆在末尾）。"""
from docx import Document

from src.readers import read_docx


def _make_docx(tmp_path):
    p = tmp_path / "order.docx"
    doc = Document()
    doc.add_heading("第一章", level=1)
    doc.add_paragraph("第一章正文")
    t1 = doc.add_table(rows=1, cols=2)
    t1.cell(0, 0).text = "T1a"
    t1.cell(0, 1).text = "T1b"
    doc.add_heading("第二章", level=1)
    doc.add_paragraph("第二章正文")
    t2 = doc.add_table(rows=1, cols=2)
    t2.cell(0, 0).text = "T2a"
    t2.cell(0, 1).text = "T2b"
    doc.save(str(p))
    return str(p)


def test_read_docx_tables_in_document_order(tmp_path):
    d = read_docx(_make_docx(tmp_path))
    kinds = [b["kind"] for b in d["blocks"]]
    assert kinds == ["heading", "paragraph", "table",
                     "heading", "paragraph", "table"]


def test_read_docx_table_rows_content(tmp_path):
    d = read_docx(_make_docx(tmp_path))
    tables = [b for b in d["blocks"] if b["kind"] == "table"]
    assert len(tables) == 2
    assert tables[0]["rows"] == [["T1a", "T1b"]]
    assert tables[1]["rows"] == [["T2a", "T2b"]]


def test_read_docx_table_follows_its_chapter(tmp_path):
    """每个表格必须出现在对应章节标题之后、下一章标题之前。"""
    d = read_docx(_make_docx(tmp_path))
    blocks = d["blocks"]
    idx_h2 = next(i for i, b in enumerate(blocks)
                  if b["kind"] == "heading" and b["text"] == "第二章")
    idx_t1 = next(i for i, b in enumerate(blocks)
                  if b["kind"] == "table" and b["rows"] == [["T1a", "T1b"]])
    assert idx_t1 < idx_h2
