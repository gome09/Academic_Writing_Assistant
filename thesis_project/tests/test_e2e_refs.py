# -*- coding: utf-8 -*-
"""参考资料模式端到端：真实文件 -> main() -> 校验生成的 docx。"""
import sys

import docx as docx_lib
import pytest

from src import main as main_mod
from src import synthesizer


@pytest.fixture()
def refs_input(tmp_path):
    d = tmp_path / "input"
    d.mkdir()
    (d / "topic.md").write_text(
        "---\nauthor: 张三\n---\n# 基于X的Y系统\n\n研究内容：……\n",
        encoding="utf-8")
    (d / "文献A.md").write_text("# 文A\n\n目标检测综述正文。\n",
                                encoding="utf-8")
    (d / "数据.csv").write_text("组别,精度\nA,0.91\n", encoding="utf-8")
    return d


def _fake_chat_json(system, user):
    if "文献调研助手" in system:
        return {"title": "文A", "authors": ["李四"], "year": "2023",
                "topic": "目标检测", "method": "综述", "conclusion": "有效",
                "quotes": ["检测器分两类"]}
    if "论文结构顾问" in system:
        return {"chapters": [
            {"title": "绪论", "kind": "intro", "cards": []},
            {"title": "相关技术综述", "kind": "review", "cards": [0]},
            {"title": "系统设计", "kind": "core", "cards": [0]},
            {"title": "总结与展望", "kind": "conclusion", "cards": []}]}
    if "学术写作助手" in system:
        return {"paras": ["综述正文[1]。"]}
    if "写作教练" in system:
        return {"系统设计": ["先画架构图"]}
    raise AssertionError("未知提示词")


def test_e2e_refs_mode_generates_word_only(refs_input, tmp_path, monkeypatch):
    monkeypatch.setenv("LLM_API_KEY", "sk-test")
    monkeypatch.delenv("LLM_VISION_MODEL", raising=False)
    monkeypatch.setattr(synthesizer, "_chat_json", _fake_chat_json)
    out = tmp_path / "output"
    monkeypatch.setattr(sys, "argv",
                        ["main.py", "--input", str(refs_input),
                         "--output", str(out)])
    assert main_mod.main() == 0

    assert (out / "论文草案.docx").exists()
    assert not (out / "答辩PPT草案.pptx").exists()     # refs 模式不产 PPT

    doc = docx_lib.Document(str(out / "论文草案.docx"))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "基于X的Y系统" in text
    assert "相关技术综述" in text
    assert "<AI生成，请核对>" in text
    assert "李四. 文A[J]. «请补全期刊», 2023." in text
    assert "某刊" not in text                 # 不接受 LLM 编造的期刊字段
    assert len(doc.tables) >= 1                        # csv 表格已插入


def test_e2e_draft_mode_untouched(tmp_path, monkeypatch):
    """无 topic 文件时走原流程，仍生成两份产物（回归保护）。"""
    d = tmp_path / "input"
    d.mkdir()
    (d / "草稿.md").write_text("# 我的论文\n\n## 绪论\n\n正文。\n",
                               encoding="utf-8")
    out = tmp_path / "output"
    monkeypatch.delenv("LLM_API_KEY", raising=False)
    monkeypatch.setattr(sys, "argv",
                        ["main.py", "--input", str(d), "--output", str(out)])
    assert main_mod.main() == 0
    assert (out / "论文草案.docx").exists()
    assert (out / "答辩PPT草案.pptx").exists()
