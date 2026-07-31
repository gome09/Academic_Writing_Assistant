# -*- coding: utf-8 -*-
"""read_docx 把 Heading 1/2/3 映射到 level 1/2/3 链路。"""
from docx import Document

from src.readers import read_docx


def _make_docx(tmp_path):
    p = tmp_path / "h.docx"
    doc = Document()
    doc.add_heading("章1", level=1)
    doc.add_heading("1.1", level=2)
    doc.add_heading("1.1.1", level=3)
    doc.add_paragraph("正文段")
    doc.save(str(p))
    return str(p)


def test_read_docx_heading_levels(tmp_path):
    d = read_docx(_make_docx(tmp_path))
    levels = [(b["kind"], b["level"], b["text"])
              for b in d["blocks"] if b["kind"] == "heading"]
    assert ("heading", 1, "章1") in levels
    assert ("heading", 2, "1.1") in levels
    assert ("heading", 3, "1.1.1") in levels


def test_read_docx_paragraphs_remain(tmp_path):
    d = read_docx(_make_docx(tmp_path))
    paras = [b["text"] for b in d["blocks"] if b["kind"] == "paragraph"]
    assert "正文段" in paras
